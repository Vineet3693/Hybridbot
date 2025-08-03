
import json
import time
from datetime import datetime
from typing import List, Dict
import streamlit as st
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
import io
from fpdf import FPDF
import markdown

class ChatExporter:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        
    def export_to_docx(self, chat_history: List[Dict], filename: str = None) -> bytes:
        """Export chat history to DOCX format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Hybrid AI Bot - Chat History', 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total conversations: {len(chat_history)}")
        doc.add_paragraph("")
        
        # Add each conversation
        for i, chat in enumerate(chat_history, 1):
            # Add conversation header
            header = doc.add_heading(f'Conversation {i}', level=1)
            
            # Add timestamp
            timestamp = datetime.fromtimestamp(chat['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
            doc.add_paragraph(f"Time: {timestamp}")
            
            # Add question
            question_para = doc.add_paragraph()
            question_run = question_para.add_run("Question: ")
            question_run.bold = True
            question_para.add_run(chat['question'])
            
            # Add answer
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run("Answer: ")
            answer_run.bold = True
            answer_para.add_run(chat['answer'])
            
            # Add sources if available
            sources = chat.get('sources', {})
            if sources.get('pdf') or sources.get('web'):
                sources_para = doc.add_paragraph()
                sources_run = sources_para.add_run("Sources: ")
                sources_run.bold = True
                
                if sources.get('pdf'):
                    doc.add_paragraph(f"• PDF sources: {len(sources['pdf'])} documents", style='List Bullet')
                    for j, pdf_source in enumerate(sources['pdf'][:3], 1):  # Show first 3 sources
                        doc.add_paragraph(f"  PDF {j}: {pdf_source[:100]}...", style='List Bullet 2')
                
                if sources.get('web'):
                    doc.add_paragraph(f"• Web sources: {len(sources['web'])} results", style='List Bullet')
                    for j, web_source in enumerate(sources['web'][:3], 1):  # Show first 3 sources
                        doc.add_paragraph(f"  Web {j}: {web_source['title']} - {web_source['snippet'][:100]}...", style='List Bullet 2')
            
            # Add separator
            doc.add_paragraph("─" * 50)
            doc.add_paragraph("")
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
    
    def export_to_pdf_reportlab(self, chat_history: List[Dict], filename: str = None) -> bytes:
        """Export chat history to PDF format using ReportLab"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=HexColor('#2E4057'),
            alignment=1  # Center
        )
        
        question_style = ParagraphStyle(
            'Question',
            parent=self.styles['Normal'],
            fontSize=12,
            textColor=HexColor('#1E88E5'),
            fontName='Helvetica-Bold',
            spaceAfter=6
        )
        
        answer_style = ParagraphStyle(
            'Answer',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leftIndent=20
        )
        
        source_style = ParagraphStyle(
            'Source',
            parent=self.styles['Normal'],
            fontSize=9,
            textColor=HexColor('#666666'),
            leftIndent=40,
            spaceAfter=6
        )
        
        # Add title
        story.append(Paragraph("Hybrid AI Bot - Chat History", title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", self.styles['Normal']))
        story.append(Paragraph(f"Total conversations: {len(chat_history)}", self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Add each conversation
        for i, chat in enumerate(chat_history, 1):
            # Conversation header
            story.append(Paragraph(f"Conversation {i}", self.styles['Heading2']))
            
            # Timestamp
            timestamp = datetime.fromtimestamp(chat['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
            story.append(Paragraph(f"Time: {timestamp}", self.styles['Normal']))
            story.append(Spacer(1, 6))
            
            # Question
            story.append(Paragraph(f"Q: {chat['question']}", question_style))
            
            # Answer
            story.append(Paragraph(f"A: {chat['answer']}", answer_style))
            
            # Sources
            sources = chat.get('sources', {})
            if sources.get('pdf') or sources.get('web'):
                story.append(Paragraph("Sources:", self.styles['Heading4']))
                
                if sources.get('pdf'):
                    story.append(Paragraph(f"PDF sources ({len(sources['pdf'])} documents):", source_style))
                    for j, pdf_source in enumerate(sources['pdf'][:3], 1):
                        story.append(Paragraph(f"• PDF {j}: {pdf_source[:150]}...", source_style))
                
                if sources.get('web'):
                    story.append(Paragraph(f"Web sources ({len(sources['web'])} results):", source_style))
                    for j, web_source in enumerate(sources['web'][:3], 1):
                        story.append(Paragraph(f"• {web_source['title']}: {web_source['snippet'][:150]}...", source_style))
            
            story.append(Spacer(1, 20))
            
            # Add page break every 3 conversations
            if i % 3 == 0 and i < len(chat_history):
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_pdf_fpdf(self, chat_history: List[Dict], filename: str = None) -> bytes:
        """Export chat history to PDF format using FPDF (alternative method)"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        
        # Title
        pdf.cell(0, 10, 'Hybrid AI Bot - Chat History', 0, 1, 'C')
        pdf.ln(10)
        
        # Metadata
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1)
        pdf.cell(0, 10, f'Total conversations: {len(chat_history)}', 0, 1)
        pdf.ln(10)
        
        # Add each conversation
        for i, chat in enumerate(chat_history, 1):
            # Check if new page is needed
            if pdf.get_y() > 250:
                pdf.add_page()
            
            # Conversation header
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, f'Conversation {i}', 0, 1)
            
            # Timestamp
            pdf.set_font('Arial', '', 9)
            timestamp = datetime.fromtimestamp(chat['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
            pdf.cell(0, 8, f'Time: {timestamp}', 0, 1)
            
            # Question
            pdf.set_font('Arial', 'B', 10)
            pdf.cell(0, 8, 'Question:', 0, 1)
            pdf.set_font('Arial', '', 10)
            
            # Handle long text by splitting into multiple lines
            question_lines = self._split_text(chat['question'], 80)
            for line in question_lines:
                pdf.cell(0, 6, line.encode('latin-1', 'replace').decode('latin-1'), 0, 1)
            
            pdf.ln(3)
            
            # Answer
            pdf.set_font('Arial', 'B', 10)
            pdf.cell(0, 8, 'Answer:', 0, 1)
            pdf.set_font('Arial', '', 10)
            
            answer_lines = self._split_text(chat['answer'], 80)
            for line in answer_lines:
                if pdf.get_y() > 270:
                    pdf.add_page()
                pdf.cell(0, 6, line.encode('latin-1', 'replace').decode('latin-1'), 0, 1)
            
            pdf.ln(5)
            
            # Sources
            sources = chat.get('sources', {})
            if sources.get('pdf') or sources.get('web'):
                pdf.set_font('Arial', 'B', 9)
                pdf.cell(0, 6, 'Sources:', 0, 1)
                pdf.set_font('Arial', '', 8)
                
                if sources.get('pdf'):
                    pdf.cell(0, 5, f'PDF sources: {len(sources["pdf"])} documents', 0, 1)
                
                if sources.get('web'):
                    pdf.cell(0, 5, f'Web sources: {len(sources["web"])} results', 0, 1)
            
            pdf.ln(8)
        
        return pdf.output(dest='S').encode('latin-1')
    
    def _split_text(self, text: str, max_chars: int) -> List[str]:
        """Split text into lines that fit within max_chars"""
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            if len(current_line + word) <= max_chars:
                current_line += word + " "
            else:
                if current_line:
                    lines.append(current_line.strip())
                current_line = word + " "
        
        if current_line:
            lines.append(current_line.strip())
        
        return lines
    
    def export_to_markdown(self, chat_history: List[Dict]) -> str:
        """Export chat history to Markdown format"""
        md_content = []
        md_content.append("# Hybrid AI Bot - Chat History\n")
        md_content.append(f"**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        md_content.append(f"**Total conversations:** {len(chat_history)}\n")
        md_content.append("---\n")
        
        for i, chat in enumerate(chat_history, 1):
            md_content.append(f"## Conversation {i}\n")
            
            timestamp = datetime.fromtimestamp(chat['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
            md_content.append(f"**Time:** {timestamp}\n")
            
            md_content.append(f"**Question:** {chat['question']}\n")
            md_content.append(f"**Answer:** {chat['answer']}\n")
            
            sources = chat.get('sources', {})
            if sources.get('pdf') or sources.get('web'):
                md_content.append("**Sources:**\n")
                if sources.get('pdf'):
                    md_content.append(f"- PDF sources: {len(sources['pdf'])} documents\n")
                if sources.get('web'):
                    md_content.append(f"- Web sources: {len(sources['web'])} results\n")
            
            md_content.append("---\n")
        
        return "\n".join(md_content)
